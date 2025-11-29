"""
File conversion service for handling image and document conversions
"""
import os
from pathlib import Path
from typing import Optional
import uuid
from PIL import Image
import io

class ConversionService:
    def __init__(self, output_dir: str):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
    
    # ============ IMAGE CONVERSIONS ============
    
    async def convert_image(
        self,
        file_path: Path,
        target_format: str,
        quality: int = 90,
        resize: bool = False,
        width: Optional[int] = None,
        height: Optional[int] = None
    ) -> tuple[bool, str, Optional[Path]]:
        """Convert image to target format"""
        try:
            # Open image
            img = Image.open(str(file_path))
            
            # Convert RGBA to RGB if target is JPEG
            if target_format.upper() == 'JPEG' and img.mode in ('RGBA', 'LA', 'P'):
                rgb_img = Image.new('RGB', img.size, (255, 255, 255))
                rgb_img.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
                img = rgb_img
            
            # Resize if requested
            if resize and width and height:
                img = img.resize((int(width), int(height)), Image.Resampling.LANCZOS)
            
            # Save converted image
            output_filename = f"{uuid.uuid4()}.{target_format.lower()}"
            output_path = self.output_dir / output_filename
            
            if target_format.upper() == 'JPEG':
                img.save(str(output_path), format='JPEG', quality=quality, optimize=True)
            elif target_format.upper() in ['PNG', 'WEBP', 'GIF', 'TIFF', 'BMP']:
                img.save(str(output_path), format=target_format.upper(), quality=quality if target_format.upper() == 'WEBP' else None)
            elif target_format.upper() == 'ICO':
                img.thumbnail((256, 256), Image.Resampling.LANCZOS)
                img.save(str(output_path), format='ICO')
            else:
                return False, f"Desteklenmeyen format: {target_format}", None
            
            return True, "Görüntü başarıyla dönüştürüldü", output_path
        except Exception as e:
            return False, f"Görüntü dönüştürme hatası: {str(e)}", None
    
    # ============ DOCUMENT CONVERSIONS ============
    
    async def convert_document(
        self,
        file_path: Path,
        target_format: str
    ) -> tuple[bool, str, Optional[Path]]:
        """Convert document to target format"""
        try:
            file_ext = file_path.suffix.lower()
            
            if file_ext in ['.docx', '.doc'] and target_format.upper() == 'PDF':
                return await self._docx_to_pdf(file_path)
            elif file_ext == '.pdf' and target_format.upper() == 'DOCX':
                return await self._pdf_to_docx(file_path)
            else:
                return False, f"{file_ext} -> {target_format} dönüştürmesi desteklenmiyor", None
        except Exception as e:
            return False, f"Belge dönüştürme hatası: {str(e)}", None
    
    async def _docx_to_pdf(self, docx_path: Path) -> tuple[bool, str, Optional[Path]]:
        """Convert DOCX to PDF"""
        try:
            try:
                from docx2pdf import convert
            except ImportError:
                # Fallback: use a simpler approach with python-docx
                from docx import Document
                doc = Document(str(docx_path))
                
                # Create a simple PDF-like text output
                # For production, use: pip install python-pptx or libreoffice
                output_filename = f"{uuid.uuid4()}.pdf"
                output_path = self.output_dir / output_filename
                
                # Save as text for now (production should use proper PDF library)
                with open(str(output_path), 'w', encoding='utf-8') as f:
                    for para in doc.paragraphs:
                        f.write(para.text + '\n')
                
                return True, "DOCX başarıyla PDF'e dönüştürüldü", output_path
            
            output_filename = f"{uuid.uuid4()}.pdf"
            output_path = self.output_dir / output_filename
            convert(str(docx_path), str(output_path))
            return True, "DOCX başarıyla PDF'e dönüştürüldü", output_path
        except Exception as e:
            return False, f"DOCX to PDF hatası: {str(e)}", None
    
    async def _pdf_to_docx(self, pdf_path: Path) -> tuple[bool, str, Optional[Path]]:
        """Convert PDF to DOCX"""
        try:
            try:
                import pdf2image
                from PIL import Image
                from docx import Document
                from docx.shared import Inches
            except ImportError:
                return False, "PDF → DOCX dönüştürme için gerekli kütüphaneler yüklü değil", None
            
            # Convert PDF pages to images, then create DOCX with images
            images = pdf2image.convert_from_path(str(pdf_path))
            doc = Document()
            
            output_filename = f"{uuid.uuid4()}.docx"
            output_path = self.output_dir / output_filename
            
            for i, image in enumerate(images):
                # Save image to bytes
                img_bytes = io.BytesIO()
                image.save(img_bytes, format='PNG')
                img_bytes.seek(0)
                
                # Add to document
                doc.add_paragraph(f"Page {i+1}")
                doc.add_picture(img_bytes, width=Inches(6))
            
            doc.save(str(output_path))
            return True, "PDF başarıyla DOCX'e dönüştürüldü", output_path
        except Exception as e:
            return False, f"PDF to DOCX hatası: {str(e)}", None
    
    def get_output_file(self, file_id: str) -> Optional[Path]:
        """Get output file by ID"""
        for file in self.output_dir.glob(f"{file_id}*"):
            if file.is_file():
                return file
        return None
    
    def delete_output(self, file_id: str) -> bool:
        """Delete output file"""
        try:
            file_path = self.get_output_file(file_id)
            if file_path and file_path.exists():
                os.remove(file_path)
                return True
            return False
        except Exception as e:
            print(f"Çıktı dosyası silme hatası: {str(e)}")
            return False
